#!/usr/bin/env python3
"""
Inkstain Trail Certificate Generator v2
Reads Word document metadata + extension Trail JSON + author note.
Produces a unified certificate of authorship.
Never reads manuscript content.
"""

import sys, os, json, hashlib, datetime
from io import BytesIO

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch

INK       = (27/255,  42/255,  59/255)
PARCHMENT = (245/255, 242/255, 235/255)
AMBER     = (200/255, 149/255, 107/255)
LIGHT     = (236/255, 232/255, 220/255)

PLATFORM_NAMES = {
    'claude':'Claude','chatgpt':'ChatGPT','gemini':'Gemini',
    'perplexity':'Perplexity','replit':'Replit','copilot':'Copilot',
    'poe':'Poe','sudowrite':'Sudowrite','novelai':'NovelAI',
    'jasper':'Jasper','writesonic':'Writesonic','rytr':'Rytr',
    'copy_ai':'Copy.ai','hyperwrite':'HyperWrite','novelcrafter':'NovelCrafter',
    'grok':'Grok','mistral':'Mistral','cohere':'Cohere',
    'huggingface':'HuggingFace','google_docs':'Google Docs','notion':'Notion',
    'character_ai':'Character.AI'
}


def extract_metadata(docx_bytes, filename="document.docx", form_author=None):
    meta = {
        "filename": filename, "created": None, "modified": None,
        "author": None, "revision_count": None,
        "total_edit_time_minutes": None, "word_count": None,
        "paragraph_count": None, "character_count": None,
        "page_count": None,
        "captured_at": datetime.datetime.utcnow().isoformat() + "Z",
        "warnings": []
    }
    try:
        doc = Document(BytesIO(docx_bytes))
        cp = doc.core_properties
        if cp.created:
            meta["created"] = cp.created.strftime("%B %d, %Y") if hasattr(cp.created,'strftime') else str(cp.created)
        if cp.modified:
            meta["modified"] = cp.modified.strftime("%B %d, %Y") if hasattr(cp.modified,'strftime') else str(cp.modified)
        doc_author = cp.author if cp.author and cp.author.strip() and cp.author.strip().lower() not in ['un-named','unnamed',''] else None
        meta["author"] = form_author or doc_author or None
        if cp.revision: meta["revision_count"] = cp.revision
        try:
            app_props = doc.part.package.part_related_by(
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties')
            if app_props:
                root = app_props._element
                ns = 'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties'
                def get_el(tag):
                    el = root.find(f'{{{ns}}}{tag}')
                    return int(el.text) if el is not None and el.text else None
                meta["word_count"] = get_el('Words')
                meta["page_count"] = get_el('Pages')
                meta["character_count"] = get_el('Characters')
                meta["total_edit_time_minutes"] = get_el('TotalTime')
                meta["paragraph_count"] = get_el('Paragraphs')
        except Exception:
            meta["warnings"].append("Extended properties unavailable")
        if not meta["word_count"]:
            total = sum(len(p.text.split()) for p in doc.paragraphs if p.text.strip())
            meta["word_count"] = total if total > 0 else None
            meta["paragraph_count"] = sum(1 for p in doc.paragraphs if p.text.strip())
            if total > 0: meta["warnings"].append("Word count estimated from document structure")
    except PackageNotFoundError:
        raise ValueError("File does not appear to be a valid Word document (.docx)")
    except Exception as e:
        raise ValueError(f"Could not read document: {str(e)}")
    return meta


def parse_trail(trail_json_str):
    """Parse extension Trail JSON into summary stats."""
    try:
        data = json.loads(trail_json_str)
        trail = data.get('trail', [])
        state = data.get('state', {})
        if not trail:
            return None

        ai_copies = [e for e in trail if e.get('type') == 'ai_copy']
        doc_pastes = [e for e in trail if e.get('type') == 'doc_paste']

        platforms = {}
        for ev in ai_copies:
            p = ev.get('platform', 'unknown')
            platforms[p] = platforms.get(p, 0) + 1

        total_chars_copied = sum(e.get('size_chars', 0) for e in ai_copies)
        total_chars_pasted = sum(e.get('size_chars', 0) for e in doc_pastes)

        # Date range
        timestamps = [e.get('timestamp') for e in trail if e.get('timestamp')]
        first_event = min(timestamps) if timestamps else None
        last_event = max(timestamps) if timestamps else None

        if first_event:
            first_event = datetime.datetime.fromisoformat(first_event.replace('Z','+00:00')).strftime("%B %d, %Y")
        if last_event:
            last_event = datetime.datetime.fromisoformat(last_event.replace('Z','+00:00')).strftime("%B %d, %Y")

        return {
            'total_events': len(trail),
            'ai_copies': len(ai_copies),
            'doc_pastes': len(doc_pastes),
            'platforms': platforms,
            'total_chars_copied': total_chars_copied,
            'total_chars_pasted': total_chars_pasted,
            'first_event': first_event,
            'last_event': last_event,
            'manuscript_title': state.get('manuscript_title', ''),
            'author_name': state.get('author_name', ''),
        }
    except Exception:
        return None


def format_edit_time(minutes):
    if not minutes: return "not recorded"
    if minutes < 60: return f"{minutes} minutes"
    hours = minutes // 60
    mins = minutes % 60
    return f"{hours}h {mins}m" if mins else f"{hours} hour{'s' if hours != 1 else ''}"


def wrap_text(text, max_chars):
    """Simple word wrapper."""
    words = text.split()
    lines = []
    current = []
    count = 0
    for word in words:
        if count + len(word) + 1 > max_chars and current:
            lines.append(' '.join(current))
            current = [word]
            count = len(word)
        else:
            current.append(word)
            count += len(word) + 1
    if current:
        lines.append(' '.join(current))
    return lines


def generate_certificate_pdf(author, title, metadata, trail_summary, author_note, disclosure_level="summary"):
    buffer = BytesIO()
    W, H = letter
    c = canvas.Canvas(buffer, pagesize=letter)

    # Background
    c.setFillColorRGB(*PARCHMENT)
    c.rect(0, 0, W, H, fill=1, stroke=0)

    # Borders
    c.setStrokeColorRGB(*INK)
    c.setLineWidth(0.8)
    c.rect(0.35*inch, 0.35*inch, W-0.7*inch, H-0.7*inch, fill=0, stroke=1)
    c.setLineWidth(0.3)
    c.rect(0.45*inch, 0.45*inch, W-0.9*inch, H-0.9*inch, fill=0, stroke=1)

    # Header band
    c.setFillColorRGB(*INK)
    c.rect(0.35*inch, H-1.8*inch, W-0.7*inch, 1.45*inch, fill=1, stroke=0)

    c.setFillColorRGB(*AMBER)
    c.setFont("Helvetica", 8)
    c.drawCentredString(W/2, H-0.85*inch, "C E R T I F I C A T E   O F   A U T H O R S H I P")

    # Inkstain wordmark — Ink in parchment, stain in amber
    ink_w = c.stringWidth("Ink", "Helvetica-Bold", 28)
    total_w = c.stringWidth("Inkstain", "Helvetica-Bold", 28)
    start_x = W/2 - total_w/2
    c.setFillColorRGB(*PARCHMENT)
    c.setFont("Helvetica-Bold", 28)
    c.drawString(start_x, H-1.38*inch, "Ink")
    c.setFillColorRGB(*AMBER)
    c.setFont("Helvetica-BoldOblique", 28)
    c.drawString(start_x + ink_w, H-1.38*inch, "stain")

    # Rule
    c.setStrokeColorRGB(*AMBER)
    c.setLineWidth(0.5)
    margin = 1.2*inch
    c.line(margin, H-1.85*inch, W-margin, H-1.85*inch)

    # Certifies
    c.setFillColorRGB(*INK)
    c.setFont("Helvetica-Oblique", 11)
    c.drawCentredString(W/2, H-2.35*inch, "This certifies that the manuscript")

    c.setFont("Helvetica-BoldOblique", 20)
    title_display = f'"{title}"' if len(title) < 44 else f'"{title[:41]}..."'
    c.drawCentredString(W/2, H-2.82*inch, title_display)

    c.setFont("Helvetica-Oblique", 11)
    c.drawCentredString(W/2, H-3.15*inch, "was composed by")

    c.setFont("Helvetica-Bold", 22)
    c.drawCentredString(W/2, H-3.6*inch, author.upper())

    # Divider
    c.setStrokeColorRGB(*INK)
    c.setLineWidth(0.3)
    c.setDash(2, 4)
    c.line(1.5*inch, H-3.9*inch, W-1.5*inch, H-3.9*inch)
    c.setDash()

    # THE TRAIL header
    c.setFont("Helvetica", 8)
    c.setFillColorRGB(*AMBER)
    c.drawCentredString(W/2, H-4.2*inch, "T H E   T R A I L")

    # Trail rows
    data_y = H-4.55*inch
    row_h = 0.28*inch
    left_col = 1.3*inch
    right_col = W-1.3*inch

    def trail_row(label, value, y, highlight=False):
        if highlight:
            c.setFillColorRGB(*LIGHT)
            c.rect(left_col-0.1*inch, y-0.05*inch, right_col-left_col+0.2*inch, row_h-0.03*inch, fill=1, stroke=0)
        c.setFillColorRGB(*INK)
        c.setFont("Helvetica", 9)
        c.setFillAlpha(0.45)
        c.drawString(left_col, y+0.07*inch, label)
        c.setFillAlpha(1.0)
        c.setFont("Helvetica-Bold", 9)
        c.drawRightString(right_col, y+0.07*inch, str(value) if value else "—")
        c.setStrokeColorRGB(*INK)
        c.setLineWidth(0.2)
        c.setStrokeAlpha(0.1)
        c.line(left_col, y-0.02*inch, right_col, y-0.02*inch)
        c.setStrokeAlpha(1.0)
        c.setFillAlpha(1.0)

    rows = []

    # Document metadata rows
    if metadata and metadata.get("word_count"):
        wc = f"{metadata['word_count']:,} words"
        if metadata.get("page_count"): wc += f"  ·  {metadata['page_count']} pages"
        rows.append(("Manuscript length", wc))
    if metadata and metadata.get("created"):
        rows.append(("Document created", metadata["created"]))
    if metadata and metadata.get("modified"):
        rows.append(("Last modified", metadata["modified"]))
    if metadata and metadata.get("total_edit_time_minutes"):
        rows.append(("Total editing time", format_edit_time(metadata["total_edit_time_minutes"])))
    if metadata and metadata.get("revision_count"):
        rows.append(("Revision count", str(metadata["revision_count"])))

    # Trail event rows — Gap 1 data
    if trail_summary:
        rows.append(("Trail recording started", trail_summary.get('first_event', '—')))
        rows.append(("AI copy events", str(trail_summary['ai_copies'])))
        rows.append(("Document paste events", str(trail_summary['doc_pastes'])))

        if trail_summary['platforms'] and disclosure_level in ['standard', 'full']:
            platform_str = ', '.join(
                f"{PLATFORM_NAMES.get(p, p)} ({n})"
                for p, n in sorted(trail_summary['platforms'].items(), key=lambda x: -x[1])
            )
            # Truncate if too long
            if len(platform_str) > 55:
                platform_str = platform_str[:52] + '...'
            rows.append(("AI platforms used", platform_str))

        if trail_summary['total_chars_copied'] > 0 and disclosure_level in ['standard', 'full']:
            rows.append(("Total chars from AI", f"{trail_summary['total_chars_copied']:,}"))

    if not rows:
        rows.append(("Trail captured", datetime.datetime.utcnow().strftime("%B %d, %Y")))

    for i, (label, value) in enumerate(rows):
        trail_row(label, value, data_y-(i*row_h), highlight=(i%2==0))

    current_y = data_y - (len(rows) * row_h) - 0.2*inch

    # AUTHOR NOTE — Gap 3
    if author_note and author_note.strip():
        c.setStrokeColorRGB(*AMBER)
        c.setLineWidth(0.3)
        c.line(left_col, current_y, right_col, current_y)
        current_y -= 0.2*inch

        c.setFont("Helvetica", 8)
        c.setFillColorRGB(*AMBER)
        c.drawString(left_col, current_y, "A U T H O R ' S   N O T E")
        current_y -= 0.2*inch

        note_lines = wrap_text(author_note.strip(), 85)
        c.setFont("Helvetica-Oblique", 9)
        c.setFillColorRGB(*INK)
        for line in note_lines[:4]:  # max 4 lines
            c.drawString(left_col, current_y, line)
            current_y -= 0.16*inch

        # Timestamp the note
        note_ts = datetime.datetime.utcnow().strftime("%B %d, %Y")
        c.setFont("Helvetica", 8)
        c.setFillAlpha(0.38)
        c.drawString(left_col, current_y, f"Author-supplied statement, recorded {note_ts}")
        c.setFillAlpha(1.0)
        current_y -= 0.25*inch

    # Capture note
    c.setFont("Helvetica-Oblique", 8)
    c.setFillColorRGB(*INK)
    c.setFillAlpha(0.4)
    capture_date = datetime.datetime.utcnow().strftime("%B %d, %Y")

    source_note = "Document metadata"
    if trail_summary and metadata: source_note = "Document metadata + extension Trail"
    elif trail_summary: source_note = "Extension Trail (no document uploaded)"

    c.drawCentredString(W/2, current_y-0.1*inch,
        f"Trail captured: {capture_date}   ·   {source_note} — manuscript text never read or stored.")
    c.drawCentredString(W/2, current_y-0.28*inch,
        "This certificate reflects recorded history. Inkstain does not verify manuscript content.")
    c.setFillAlpha(1.0)

    # Hash
    hash_input = f"{author}:{title}:{metadata.get('created','') if metadata else ''}:{trail_summary['ai_copies'] if trail_summary else 0}:{capture_date}"
    doc_hash = hashlib.sha256(hash_input.encode()).hexdigest()
    c.setFont("Helvetica", 7)
    c.setFillAlpha(0.25)
    c.drawCentredString(W/2, current_y-0.46*inch, f"verification hash  ·  {doc_hash[:48]}")
    c.setFillAlpha(1.0)

    # Bottom
    c.setStrokeColorRGB(*AMBER)
    c.setLineWidth(0.5)
    c.line(margin, 0.9*inch, W-margin, 0.9*inch)

    c.setFont("Helvetica-Oblique", 10)
    c.setFillColorRGB(*INK)
    c.drawCentredString(W/2, 0.65*inch, "The written word will prevail.")

    c.setFont("Helvetica", 8)
    c.setFillAlpha(0.35)
    c.drawCentredString(W/2, 0.48*inch, "inkstain.ai   ·   AI Provenance for Authors")
    c.setFillAlpha(1.0)

    # Print hash to stdout so server.js can capture it
    print(f"INKSTAIN_HASH:{doc_hash}")

    c.save()
    return buffer.getvalue()


if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Usage: python3 certificate.py <file.docx|none> <author> <title> [disclosure] [trail.json] [note]")
        sys.exit(1)

    filepath = sys.argv[1]
    author = sys.argv[2]
    title = sys.argv[3]
    disclosure = sys.argv[4] if len(sys.argv) > 4 else "summary"
    trail_file = sys.argv[5] if len(sys.argv) > 5 else None
    note = sys.argv[6] if len(sys.argv) > 6 else ""

    metadata = None
    if filepath and filepath != 'none' and os.path.exists(filepath):
        with open(filepath, 'rb') as f:
            metadata = extract_metadata(f.read(), os.path.basename(filepath), form_author=author)
        print("Metadata:", json.dumps(metadata, indent=2))

    trail_summary = None
    if trail_file and os.path.exists(trail_file):
        with open(trail_file, 'r') as f:
            trail_summary = parse_trail(f.read())
        print("Trail summary:", json.dumps(trail_summary, indent=2))

    pdf_bytes = generate_certificate_pdf(author, title, metadata, trail_summary, note, disclosure)

    output_path = f"{title.replace(' ','_')}_trail_certificate.pdf"
    with open(output_path, 'wb') as f:
        f.write(pdf_bytes)
    print(f"Certificate saved to: {output_path}")
